#!/usr/bin/env python3
"""Build the client-facing DOCX review copy for the Downs Style fabric trilogy."""

from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
CHARLES = ROOT / "docs" / "continuations" / "charles"
OUTPUT = CHARLES / "downs-style-three-new-blogs-review-copy.docx"

ARTICLES = (
    CHARLES / "cotton-summer-post.md",
    CHARLES / "silk-summer-post.md",
    CHARLES / "cotton-vs-silk-post.md",
)

INDIVIDUAL_OUTPUTS = {
    ARTICLES[0]: CHARLES / "downs-style-cotton-summer-charles-review.docx",
    ARTICLES[1]: CHARLES / "downs-style-silk-charles-review.docx",
    ARTICLES[2]: CHARLES / "downs-style-cotton-vs-silk-charles-review.docx",
}

# Resolved preset: google_docs_default.
# Named overrides used consistently:
# - Article Section: centered 13 pt bold for headings embedded in Squarespace copy.
# - Metadata: muted 9.5 pt review-only article settings.
# - Popup Trigger: compact 12.5 pt bold for the client-only callout list.
# - Popup Answer: 10.5 pt body with a small left indent for callout copy.
BLACK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)


def set_run_font(run, *, size: float | None = None, color=BLACK, bold=None, italic=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style(style, *, size, color, before, after, line, bold=False):
    style.font.name = "Arial"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = style.name.startswith("Heading")


def configure_document(document: Document, *, title: str, subject: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    # 0.492 in resolves to the nearest OOXML unit: 708 twips.
    section.header_distance = Twips(708)
    section.footer_distance = Twips(708)

    styles = document.styles
    configure_style(
        styles["Normal"],
        size=11,
        color=BLACK,
        before=0,
        after=8,
        line=1.15,
    )
    configure_style(
        styles["Heading 1"],
        size=20,
        color=BLACK,
        before=20,
        after=6,
        line=1.15,
        bold=False,
    )
    configure_style(
        styles["Heading 2"],
        size=16,
        color=BLACK,
        before=18,
        after=6,
        line=1.15,
        bold=False,
    )
    configure_style(
        styles["Heading 3"],
        size=14,
        color=RGBColor(0x43, 0x43, 0x43),
        before=16,
        after=4,
        line=1.15,
        bold=False,
    )

    article_section = styles.add_style("Article Section", WD_STYLE_TYPE.PARAGRAPH)
    configure_style(
        article_section,
        size=13,
        color=BLACK,
        before=14,
        after=6,
        line=1.15,
        bold=True,
    )
    article_section.paragraph_format.keep_with_next = True

    metadata = styles.add_style("Metadata", WD_STYLE_TYPE.PARAGRAPH)
    configure_style(
        metadata,
        size=9.5,
        color=MUTED,
        before=0,
        after=3,
        line=1.15,
    )

    popup_trigger = styles.add_style("Popup Trigger", WD_STYLE_TYPE.PARAGRAPH)
    configure_style(
        popup_trigger,
        size=12.5,
        color=BLACK,
        before=10,
        after=3,
        line=1.1,
        bold=True,
    )
    popup_trigger.paragraph_format.keep_with_next = True

    popup = styles.add_style("Popup Answer", WD_STYLE_TYPE.PARAGRAPH)
    configure_style(
        popup,
        size=10.5,
        color=BLACK,
        before=0,
        after=8,
        line=1.1,
    )
    popup.paragraph_format.left_indent = Inches(0.25)

    document.core_properties.title = title
    document.core_properties.subject = subject
    document.core_properties.author = "Downs Style editorial team"
    document.core_properties.last_modified_by = "Downs Style editorial team"


def section(text: str, heading_pattern: str, next_heading: str = r"^## ") -> str:
    match = re.search(
        rf"^{heading_pattern}\s*$\n(.*?)(?={next_heading}|\Z)",
        text,
        flags=re.MULTILINE | re.DOTALL,
    )
    if not match:
        raise ValueError(f"missing section matching {heading_pattern!r}")
    return match.group(1).strip()


def metadata(text: str) -> dict[str, str]:
    settings = section(text, r"## Page settings")
    values: dict[str, str] = {}
    for line in settings.splitlines():
        match = re.match(r"- \*\*(.+?):\*\*\s*(.+)$", line)
        if match:
            values[match.group(1)] = match.group(2).strip().strip("`")
    return values


def add_inline_markdown(paragraph, text: str) -> None:
    cursor = 0
    for match in re.finditer(r"(\*\*.+?\*\*|\*[^*]+?\*)", text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor : match.start()]))
        token = match.group(0)  # allow-secret: Markdown token, never a credential.
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), bold=True)
        else:
            set_run_font(paragraph.add_run(token[1:-1]), italic=True)
        cursor = match.end()
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]))


def add_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(72)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run("Downs Style Summer Fabric Trilogy")
    set_run_font(run, size=26, bold=False)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(24)
    set_run_font(subtitle.add_run("Charles review copy · August 2, 2026"), color=MUTED)

    intro = document.add_paragraph()
    add_inline_markdown(
        intro,
        "Three connected articles in Charles’s voice: cotton, silk, and the garment-by-garment decision. Each includes interactive common-error callouts checked against current sources.",
    )

    for number, title in enumerate(
        (
            "Cotton Tops and Bottoms for Summer: What Kind Are You Actually Wearing?",
            "Silk, Satin, and Charmeuse: What Are You Actually Buying?",
            "Cotton vs. Silk: Who Wins Each Part of the Wardrobe?",
        ),
        start=1,
    ):
        item = document.add_paragraph()
        item.paragraph_format.space_before = Pt(8 if number == 1 else 0)
        set_run_font(item.add_run(f"{number}. "), bold=True)
        set_run_font(item.add_run(title))

    status = document.add_paragraph()
    status.paragraph_format.space_before = Pt(18)
    set_run_font(status.add_run("STATUS  "), size=9.5, color=MUTED, bold=True)
    set_run_font(
        status.add_run("Review copy. No website has been edited or published."),
        size=9.5,
        color=MUTED,
    )

    handoff = document.add_paragraph(style="Metadata")
    set_run_font(handoff.add_run("EDITORIAL HANDOFF  "), size=9.5, color=MUTED, bold=True)
    set_run_font(
        handoff.add_run(
            "Image placement, guardrails, fact-check notes, and full source lists remain in the accompanying production files."
        ),
        size=9.5,
        color=MUTED,
    )


def add_article_body(document: Document, body: str) -> None:
    centered = re.compile(
        r'^<p style="text-align: center;"><strong>(.+?)</strong></p>$'
    )
    for block in re.split(r"\n\s*\n", body):
        block = block.strip()
        if not block:
            continue
        heading = centered.match(block)
        if heading:
            paragraph = document.add_paragraph(style="Article Section")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph.add_run(heading.group(1)), size=13, bold=True)
            continue
        paragraph = document.add_paragraph()
        add_inline_markdown(paragraph, block)


def popup_blocks(text: str) -> list[tuple[str, str, str]]:
    popup_section = section(text, r"## (?:Common-error|Myth and shopping-error) pop-ups")
    blocks = re.split(r"(?=^### )", popup_section, flags=re.MULTILINE)
    popups: list[tuple[str, str, str]] = []
    for block in blocks:
        heading = re.search(r"^### (.+)$", block, flags=re.MULTILINE)
        trigger = re.search(r"^\*\*Trigger:\*\*\s*(.+)$", block, flags=re.MULTILINE)
        answer = re.search(
            r"^\*\*Pop-up:\*\*\s*(.+?)(?=\n\n|\Z)",
            block,
            flags=re.MULTILINE | re.DOTALL,
        )
        if heading and trigger and answer:
            popups.append(
                (
                    heading.group(1).strip(),
                    trigger.group(1).strip(),
                    answer.group(1).strip(),
                )
            )
    return popups


def add_popups(document: Document, text: str) -> None:
    document.add_paragraph("Interactive common-error copy", style="Heading 2")
    note = document.add_paragraph(style="Metadata")
    set_run_font(
        note.add_run(
            "Place each trigger beside the matching article section. The body remains complete when every pop-up is closed."
        ),
        size=9.5,
        color=MUTED,
        italic=True,
    )
    for _, trigger, answer in popup_blocks(text):
        heading = document.add_paragraph(style="Popup Trigger")
        add_inline_markdown(heading, trigger)
        paragraph = document.add_paragraph(style="Popup Answer")
        add_inline_markdown(paragraph, answer)


def add_article(document: Document, path: Path, *, page_break: bool = True) -> None:
    text = path.read_text(encoding="utf-8")
    settings = metadata(text)
    if page_break:
        document.add_page_break()
    document.add_paragraph(settings["Title"], style="Heading 1")

    for label in ("Suggested slug", "SEO description", "Keywords/tags"):
        paragraph = document.add_paragraph(style="Metadata")
        set_run_font(paragraph.add_run(f"{label}: "), size=9.5, color=MUTED, bold=True)
        set_run_font(paragraph.add_run(settings[label]), size=9.5, color=MUTED)

    document.add_paragraph("Article copy", style="Heading 2")
    add_article_body(document, section(text, r"## Squarespace body copy"))
    add_popups(document, text)


def audit(document: Document) -> None:
    section0 = document.sections[0]
    assert section0.page_width == Inches(8.5)
    assert section0.page_height == Inches(11)
    assert section0.top_margin == Inches(1)
    assert section0.right_margin == Inches(1)
    assert section0.bottom_margin == Inches(1)
    assert section0.left_margin == Inches(1)
    assert section0.header_distance == Twips(708)
    assert section0.footer_distance == Twips(708)

    normal = document.styles["Normal"]
    assert normal.font.name == "Arial"
    assert normal.font.size == Pt(11)
    assert normal.paragraph_format.space_after == Pt(8)
    assert normal.paragraph_format.line_spacing == 1.15

    assert all(paragraph.style.name != "Title" for paragraph in document.paragraphs)


def save_and_audit(document: Document, output: Path) -> None:
    audit(document)
    document.save(output)
    with ZipFile(output) as archive:
        assert not any(
            re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
            for name in archive.namelist()
        )
    print(output)


def build_combined() -> None:
    document = Document()
    configure_document(
        document,
        title="Downs Style Summer Fabric Trilogy",
        subject="Charles review copy: cotton, silk, and cotton versus silk",
    )
    add_cover(document)
    for article in ARTICLES:
        add_article(document, article)
    save_and_audit(document, OUTPUT)


def build_individual(article: Path, output: Path) -> None:
    text = article.read_text(encoding="utf-8")
    settings = metadata(text)
    document = Document()
    configure_document(
        document,
        title=settings["Title"],
        subject="Downs Style article review copy for Charles",
    )
    add_article(document, article, page_break=False)
    save_and_audit(document, output)


def main() -> None:
    build_combined()
    for article, output in INDIVIDUAL_OUTPUTS.items():
        build_individual(article, output)


if __name__ == "__main__":
    main()
